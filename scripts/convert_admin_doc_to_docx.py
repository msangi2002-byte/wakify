#!/usr/bin/env python3
"""
Convert ADMIN_PANEL_FEATURES_SPECIFICATION.md to DOCX
Run: pip install python-docx
Then: python scripts/convert_admin_doc_to_docx.py
"""
import re
import os

def convert_md_to_docx():
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("Please install: pip install python-docx")
        return False

    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    md_path = os.path.join(project_root, "docs", "ADMIN_PANEL_FEATURES_SPECIFICATION.md")
    docx_path = os.path.join(project_root, "docs", "ADMIN_PANEL_FEATURES_SPECIFICATION.docx")

    if not os.path.exists(md_path):
        print(f"File not found: {md_path}")
        return False

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Calibri'

    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # H1
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=0)
        # H2
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=1)
        # H3
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=2)
        # H4
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=3)
        # Horizontal rule
        elif stripped == '---':
            doc.add_paragraph('─' * 50)
        # Table row (| ... |)
        elif stripped.startswith('|') and '|' in stripped[1:]:
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if cells:
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Table Grid'
                row = table.rows[0]
                for j, cell in enumerate(cells):
                    if j < len(row.cells):
                        row.cells[j].text = cell
                i += 1
                # Add more rows if following lines are also table rows
                while i < len(lines) and lines[i].strip().startswith('|') and not lines[i].strip().replace('-','').replace('|','').strip() == '':
                    next_line = lines[i].strip()
                    if not re.match(r'^\|[\s\-:]+\|', next_line):
                        cells = [c.strip() for c in next_line.split('|')[1:-1]]
                        row = table.add_row()
                        for j, cell in enumerate(cells):
                            if j < len(row.cells):
                                row.cells[j].text = cell
                    i += 1
                i -= 1  # adjust for loop
        # Bullet list
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            # Bold for **text**
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(text)
        # Numbered (digit. )
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            doc.add_paragraph(text, style='List Number')
        # Empty
        elif not stripped:
            pass
        # Normal paragraph
        else:
            doc.add_paragraph(stripped)

        i += 1

    doc.save(docx_path)
    print(f"Created: {docx_path}")
    return True

if __name__ == "__main__":
    convert_md_to_docx()
